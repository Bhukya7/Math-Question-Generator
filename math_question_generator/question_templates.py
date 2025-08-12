from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class MultipleChoiceQuestion:
    def __init__(self, title, description, question, instruction, difficulty, order,
                 options, correct_index, explanation, subject, unit, topic, plusmarks=1):
        self.title = title
        self.description = description
        self.question = question
        self.instruction = instruction
        self.difficulty = difficulty
        self.order = order
        self.options = options
        self.correct_index = correct_index
        self.explanation = explanation
        self.subject = subject
        self.unit = unit
        self.topic = topic
        self.plusmarks = plusmarks

    def add_to_document(self, doc):
        """Add this question to a Word document"""
        # Add question heading
        heading = doc.add_paragraph(style='Heading1')
        heading_run = heading.add_run(f"Question {self.order}: {self.title}")
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)

        # Add description
        if self.description:
            desc = doc.add_paragraph(self.description, style='Intense Quote')
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add question content
        self._add_question_content(doc)

        # Add instruction
        instruction_para = doc.add_paragraph()
        instruction_para.add_run("Instruction: ").bold = True
        instruction_para.add_run(self.instruction)

        # Add options
        options_heading = doc.add_paragraph("Options:", style='Heading2')
        options_heading.runs[0].bold = True
        
        for i, option in enumerate(self.options):
            prefix = "✓" if i == self.correct_index else "○"
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(f" {prefix} {option}")

        # Add metadata
        doc.add_paragraph().add_run("Difficulty:").bold = True
        doc.add_paragraph(self.difficulty)
        
        doc.add_paragraph().add_run("Explanation:").bold = True
        doc.add_paragraph(self.explanation)
        
        doc.add_paragraph().add_run("Subject:").bold = True
        doc.add_paragraph(f"{self.subject} > {self.unit} > {self.topic}")
        
        doc.add_paragraph().add_run("Marks:").bold = True
        doc.add_paragraph(str(self.plusmarks))

        doc.add_page_break()

    def _add_question_content(self, doc):
        """Add the main question content including tables and images"""
        question_paragraphs = self.question.split('\n')
        in_table = False
        table_rows = []

        for para in question_paragraphs:
            if not para.strip():
                continue

            if para.startswith('##'):
                heading = doc.add_heading(para[2:].strip(), level=2)
                heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            elif para.startswith('|'):
                if not in_table:
                    in_table = True
                table_rows.append(para)
            else:
                if in_table:
                    self._add_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                if "[IMAGE:" in para:
                    self._add_image(doc, para)
                else:
                    doc.add_paragraph(para)

        if in_table:
            self._add_table(doc, table_rows)

    def _add_table(self, doc, rows):
        """Add properly formatted table"""
        # Filter out separator lines
        rows = [row for row in rows if not all(c in ['|', '-', ' ', ':'] for c in row)]
        if not rows:
            return

        cols = rows[0].count('|') - 1
        table = doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells = [cell.strip() for cell in row.split('|')[1:-1]]
            for j, cell in enumerate(cells):
                table.cell(i, j).text = cell
                table.cell(i, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_image(self, doc, image_text):
        """Add image with proper sizing"""
        start_idx = image_text.find("[IMAGE:") + 7
        end_idx = image_text.find("]", start_idx)
        image_path = image_text[start_idx:end_idx].strip()

        try:
            if os.path.exists(image_path):
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(image_path, width=Inches(4.5))
            else:
                doc.add_paragraph("[Image not found]").alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Image error: {e}]").alignment = WD_ALIGN_PARAGRAPH.CENTER

def init_document_styles(doc):
    """Initialize custom styles for the document"""
    styles = doc.styles

    # Custom heading style
    if 'QuestionTitle' not in styles:
        heading_style = styles.add_style('QuestionTitle', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)